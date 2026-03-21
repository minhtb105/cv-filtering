"""
PDF to DOCX converter with TopCV template styling
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional
import pdfplumber
from .template_manager import TemplateConfig, TemplateStyle


class PDFToDocxConverter:
    """Converts PDF CVs to DOCX with TopCV template styling"""
    
    def __init__(self):
        self.template_manager = None
    
    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> str:
        """Extract text content from PDF"""
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"Failed to extract PDF text from {pdf_path}: {str(e)}")
    
    @staticmethod
    def _hex_to_rgb(hex_color: str) -> tuple:
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    @staticmethod
    def _add_styled_heading(doc: Document, text: str, template: TemplateConfig, level: int = 2):
        """Add styled heading to document"""
        paragraph = doc.add_paragraph(text, style=f'Heading {level}')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Apply template styling
        for run in paragraph.runs:
            run.font.name = template.font_header
            run.font.size = Pt(14 if level == 1 else 12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*PDFToDocxConverter._hex_to_rgb(template.color_primary))
        
        return paragraph
    
    @staticmethod
    def _add_styled_text(doc: Document, text: str, template: TemplateConfig, is_bold: bool = False):
        """Add styled paragraph text"""
        paragraph = doc.add_paragraph(text)
        for run in paragraph.runs:
            run.font.name = template.font_main
            run.font.size = Pt(11)
            run.font.bold = is_bold
            if is_bold:
                run.font.color.rgb = RGBColor(*PDFToDocxConverter._hex_to_rgb(template.color_secondary))
        return paragraph
    
    @staticmethod
    def pdf_to_docx_with_template(
        pdf_path: str,
        output_path: str,
        template: TemplateConfig,
        candidate_info: Optional[dict] = None
    ) -> bool:
        """
        Convert PDF to DOCX with template styling
        
        Args:
            pdf_path: Path to source PDF
            output_path: Path to output DOCX
            template: Template configuration to apply
            candidate_info: Optional extracted candidate information
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Extract text from PDF
            pdf_text = PDFToDocxConverter.extract_text_from_pdf(pdf_path)
            
            # Create new DOCX document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = template.font_main
            font.size = Pt(11)
            
            # Add title/header with template color
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run("CURRICULUM VITAE")
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = template.font_header
            run.font.color.rgb = RGBColor(*PDFToDocxConverter._hex_to_rgb(template.color_primary))
            
            # Add template info
            if template.has_color_blocks:
                template_para = doc.add_paragraph()
                template_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = template_para.add_run(f"[{template.name} Template]")
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(*PDFToDocxConverter._hex_to_rgb(template.color_secondary))
            
            # Add spacing
            doc.add_paragraph()
            
            # Add the PDF content as structured text
            PDFToDocxConverter._add_styled_heading(doc, "PROFILE INFORMATION", template, 1)
            doc.add_paragraph(pdf_text[:500] if len(pdf_text) > 500 else pdf_text)
            
            # Parse and organize content sections
            sections_data = PDFToDocxConverter._parse_cv_sections(pdf_text)
            
            for section_name, section_content in sections_data.items():
                PDFToDocxConverter._add_styled_heading(doc, section_name, template, 2)
                doc.add_paragraph(section_content)
            
            # Add footer with template information
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer_para.add_run(f"Created with {template.name} Template")
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error converting PDF to DOCX: {str(e)}")
            return False
    
    @staticmethod
    def _parse_cv_sections(text: str) -> dict:
        """Parse CV text into sections"""
        sections = {}
        
        # Common CV section keywords
        section_keywords = [
            'EXPERIENCE', 'EDUCATION', 'SKILLS', 'LANGUAGES',
            'CERTIFICATIONS', 'PROJECTS', 'AWARDS', 'PUBLICATIONS',
            'SUMMARY', 'OBJECTIVE', 'PROFESSIONAL'
        ]
        
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line_upper = line.upper().strip()
            # Check if line is a section header
            if any(keyword in line_upper for keyword in section_keywords):
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                current_section = line.strip()
                current_content = []
            elif current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections if sections else {'CONTENT': text}
