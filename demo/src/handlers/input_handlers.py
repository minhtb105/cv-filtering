"""
Input handlers for different CV file formats
Extensible design to support PDF, DOCX, Images (OCR in v2.0)
"""
from abc import ABC, abstractmethod
from pathlib import Path
import pdfplumber
from docx import Document
from typing import Optional


class InputHandler(ABC):
    """Base class for all input handlers"""
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract raw text from file"""
        pass
    
    @abstractmethod
    def validate(self, file_path: str) -> bool:
        """Validate file format"""
        pass
    
    @property
    @abstractmethod
    def supported_formats(self) -> list:
        """Return list of supported file extensions"""
        pass


class PDFHandler(InputHandler):
    """Handle PDF CV extraction"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from PDF {file_path}: {str(e)}")
    
    def validate(self, file_path: str) -> bool:
        """Check if file is valid PDF"""
        try:
            with pdfplumber.open(file_path) as pdf:
                return len(pdf.pages) > 0
        except:
            return False
    
    @property
    def supported_formats(self) -> list:
        return [".pdf"]


class DOCXHandler(InputHandler):
    """Handle DOCX CV extraction"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX {file_path}: {str(e)}")
    
    def validate(self, file_path: str) -> bool:
        """Check if file is valid DOCX"""
        try:
            Document(file_path)
            return True
        except:
            return False
    
    @property
    def supported_formats(self) -> list:
        return [".docx", ".doc"]


class ImageOCRHandler(InputHandler):
    """Handle image-based CV extraction (stub for v2.0)"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from image using OCR (placeholder)"""
        raise NotImplementedError(
            "Image OCR support coming in v2.0. "
            "Requires pytesseract and image preprocessing."
        )
    
    def validate(self, file_path: str) -> bool:
        """Check if file is valid image"""
        # Stub validation
        path = Path(file_path)
        return path.suffix.lower() in self.supported_formats
    
    @property
    def supported_formats(self) -> list:
        return [".png", ".jpg", ".jpeg", ".jpg"]


class InputHandlerFactory:
    """Factory for creating appropriate handlers based on file type"""
    
    _handlers = {
        ".pdf": PDFHandler(),
        ".docx": DOCXHandler(),
        ".doc": DOCXHandler(),
        ".png": ImageOCRHandler(),
        ".jpg": ImageOCRHandler(),
        ".jpeg": ImageOCRHandler(),
    }
    
    @classmethod
    def get_handler(cls, file_path: str) -> InputHandler:
        """Get appropriate handler for file"""
        ext = Path(file_path).suffix.lower()
        if ext not in cls._handlers:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {list(cls._handlers.keys())}")
        return cls._handlers[ext]
    
    @classmethod
    def register_handler(cls, extension: str, handler: InputHandler):
        """Register new file format handler (for extensibility)"""
        cls._handlers[extension.lower()] = handler
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if file format is supported"""
        ext = Path(file_path).suffix.lower()
        return ext in cls._handlers


# Test
if __name__ == "__main__":
    # Test factory
    pdf_handler = InputHandlerFactory.get_handler("test.pdf")
    print(f"PDF handler: {pdf_handler.__class__.__name__}")
    
    docx_handler = InputHandlerFactory.get_handler("test.docx")
    print(f"DOCX handler: {docx_handler.__class__.__name__}")
    
    print(f"Supported formats: {list(InputHandlerFactory._handlers.keys())}")
